#!/usr/bin/env python
# -*- coding: utf-8 -*-

import os
from os.path import expanduser
import sys
import re
from docx import Document
from docx import oxml
from docx.shared import Inches, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT

home = expanduser("~")
os.chdir("%s/libreoffice" % home)
os.system("rm -f fichas.docx")

document = Document('template.docx')

def to_unicode(text):
    return unicode(text, 'utf-8')

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def header():
    table = document.add_table(1, 2)
    cells = table.rows[0].cells

    delete_paragraph(cells[0].paragraphs[0])
    p0 = cells[0].add_paragraph('',
        style=document.styles['CabeceraFichaVerde'])

    delete_paragraph(cells[1].paragraphs[0])
    p1 = cells[1].add_paragraph('',
        style=document.styles['CabeceraFichaVerde'])

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    table.columns[0].width = Cm(13.31)
    table.columns[1].width = Cm(2.09)

    text = "II PLAN DE DESARROLLO SOSTENIBLE DEL PARQUE NATURAL\n" \
           "SIERRA DE BAZA\n" \
           "Fichas de Medidas"

    run = p0.add_run(text)

    run = p1.add_run('')
    run.add_picture("logo.png", width=Cm(1.69), height=Cm(1.06))


def create_table(section):

    title = re.sub('\n*', '', section[0])

    document.add_paragraph()

    table = document.add_table(1, 1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False

    cells = table.rows[0].cells
    delete_paragraph(cells[0].paragraphs[0])

    green5 = oxml.parse_xml(r'<w:shd {} w:fill="008000"/>'.format(oxml.ns.nsdecls('w')))
    cells[0]._tc.get_or_add_tcPr().append(green5)

    p = cells[0].add_paragraph('',
        style=document.styles['CabeceraFicha'])
    run = p.add_run(to_unicode(title))

    cells = table.add_row().cells
    delete_paragraph(cells[0].paragraphs[0])

    for line in [line for line in section[1].split('\n') if line]:
        line = re.sub('^\n*', '', line)
        line = re.sub('(\d+)((\.-)?\s+)', r'\1.- ', line, 1)
        p = cells[0].add_paragraph('',
            style=document.styles['ContenidoFicha'])
        if 'DIGO MEDIDA' in title:
            run = p.add_run(to_unicode(line.upper())).bold = True
        else:
            run = p.add_run(to_unicode(line))

def process():

    delete_paragraph(document.paragraphs[0])

    data = open("medidas.txt").read()
    fichas = data.split("CÓDIGO MEDIDA")

    for ficha in [ficha for ficha in fichas if len(ficha) > 20]:

        #header()

        ficha = 'CÓDIGO MEDIDA\n' + ficha
        current_section = [None, '']
        sections = []
        for line in ficha.split('\n'):
            if line.isupper():
                if current_section[0] is not None:
                    sections.append(current_section)
                    current_section = [None, '']
                current_section[0] = line
            else:
                current_section[1] += line + '\n'
        # add last section
        sections.append(current_section)

        for section in sections:
            create_table(section)

        document.add_page_break()

if __name__ =="__main__":
    process()
    document.save("fichas.docx")
